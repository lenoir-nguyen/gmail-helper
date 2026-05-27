"""
Document service — analyze structure and populate Word/Excel files.

Supports:
  - .xlsx  (Excel, via openpyxl)
  - .docx  (Word, via python-docx)

Key design:
  - Finds the RIGHT table/sheet by fuzzy-matching agent column names against
    document headers — works on complex multi-table documents.
  - Never destroys the original document; only appends rows to existing tables.
"""

import re
from copy import deepcopy
from io import BytesIO
from typing import Any

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl.styles import Font, PatternFill, Alignment

# Word XML namespace
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─── Normalization & matching ──────────────────────────────────────────────────

def _normalize(text: str) -> str:
    """Lowercase, strip punctuation/underscores, collapse whitespace."""
    text = text.lower()
    text = re.sub(r"[_\-/&,;:()'\"]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _match_score(col: str, header: str) -> int:
    """
    Return a match score (0-3) between an agent field name and a document header.
    Higher = better match.
    """
    nc = _normalize(col)
    nh = _normalize(header)
    if not nc or not nh:
        return 0
    if nc == nh:
        return 3                          # Exact match
    if nc in nh or nh in nc:
        return 2                          # One contains the other
    # Word-level overlap
    nc_words = set(nc.split())
    nh_words = set(nh.split())
    overlap = nc_words & nh_words
    # Ignore trivial stop words
    overlap -= {"of", "and", "the", "a", "an", "in", "to", "for", "or"}
    if overlap:
        return 1
    return 0


def _best_col_for_header(header: str, columns: list[str]) -> str | None:
    """
    Given a document header, return the agent column name that best matches it.
    Returns None if no match is found.
    """
    scores = [(col, _match_score(col, header)) for col in columns]
    best_col, best_score = max(scores, key=lambda x: x[1])
    return best_col if best_score > 0 else None


# ─── Analyze ──────────────────────────────────────────────────────────────────


def analyze_document(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Read the uploaded document and return its structure.
    This is passed to the AI agent so it understands where to put data.
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext in ("xlsx", "xls"):
        return _analyze_excel(file_bytes)
    elif ext in ("docx", "doc"):
        return _analyze_word(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Use .xlsx or .docx.")


def _analyze_excel(file_bytes: bytes) -> dict:
    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True)
    sheets = []
    for ws in wb.worksheets:
        headers = [str(cell.value) for cell in ws[1] if cell.value is not None]
        sheets.append({
            "name": ws.title,
            "headers": headers,
            "row_count": ws.max_row,
        })
    return {"type": "excel", "sheets": sheets}


def _analyze_word(file_bytes: bytes) -> dict:
    doc = Document(BytesIO(file_bytes))
    tables = []
    for i, table in enumerate(doc.tables):
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            tables.append({
                "index": i,
                "headers": headers,
                "row_count": len(table.rows),
            })

    # Only send a small sample of text (context budget)
    sample_text = [
        p.text.strip()[:200] for p in doc.paragraphs if p.text.strip()
    ][:3]

    return {"type": "word", "tables": tables, "sample_text": sample_text}


# ─── Populate ─────────────────────────────────────────────────────────────────


def populate_document(file_bytes: bytes, filename: str, agent_result: dict) -> bytes:
    """
    Populate the uploaded document with the agent's extracted data.
    Returns the modified file as bytes.
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext in ("xlsx", "xls"):
        return _populate_excel(file_bytes, agent_result)
    elif ext in ("docx", "doc"):
        return _populate_word(file_bytes, agent_result)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")


# ─── Excel ────────────────────────────────────────────────────────────────────

_HEADER_FILL = PatternFill("solid", fgColor="4472C4")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
_ALT_FILL    = PatternFill("solid", fgColor="DCE6F1")


def _find_best_sheet(wb, columns: list[str]):
    """Return the worksheet whose row-1 headers best match the agent columns."""
    best_ws    = wb.active
    best_score = -1

    for ws in wb.worksheets:
        headers = [
            str(ws.cell(1, c).value or "").strip()
            for c in range(1, (ws.max_column or 1) + 2)
        ]
        score = sum(
            max((_match_score(col, h) for h in headers), default=0)
            for col in columns
        )
        if score > best_score:
            best_score = score
            best_ws = ws

    return best_ws


def _populate_excel(file_bytes: bytes, agent_result: dict) -> bytes:
    data: list[dict] = agent_result.get("data", [])
    columns: list[str] = agent_result.get("column_order", []) or list(data[0].keys())
    summary: str = agent_result.get("summary", "")

    if not data:
        raise ValueError("No data was extracted from Gmail.")

    wb = openpyxl.load_workbook(BytesIO(file_bytes))
    ws = _find_best_sheet(wb, columns)

    # Read existing headers from row 1
    existing_headers = [
        str(ws.cell(1, c).value or "").strip()
        for c in range(1, (ws.max_column or 1) + 1)
    ]
    existing_headers = [h for h in existing_headers if h]  # drop blanks

    if existing_headers:
        # Build a mapping: col_index → agent field name (fuzzy matched)
        col_map: dict[int, str] = {}
        for col_idx, header in enumerate(existing_headers, 1):
            field = _best_col_for_header(header, columns)
            if field:
                col_map[col_idx] = field

        # Find the first truly empty data row (skip header)
        start_row = ws.max_row + 1
        # Walk back to find the last non-empty row so we don't leave gaps
        for r in range(ws.max_row, 1, -1):
            row_vals = [ws.cell(r, c).value for c in range(1, len(existing_headers) + 1)]
            if any(v for v in row_vals):
                start_row = r + 1
                break

        for row_data in data:
            for col_idx, field in col_map.items():
                value = row_data.get(field, "")
                cell = ws.cell(start_row, col_idx, str(value) if value is not None else "")
                if start_row % 2 == 0:
                    cell.fill = _ALT_FILL
            start_row += 1

    else:
        # Blank sheet — write headers + data from scratch
        for col_idx, header in enumerate(columns, 1):
            cell = ws.cell(1, col_idx, header)
            cell.fill = _HEADER_FILL
            cell.font = _HEADER_FONT
            cell.alignment = Alignment(horizontal="center")

        for row_idx, row_data in enumerate(data, 2):
            for col_idx, field in enumerate(columns, 1):
                value = row_data.get(field, "")
                cell = ws.cell(row_idx, col_idx, str(value) if value is not None else "")
                if row_idx % 2 == 0:
                    cell.fill = _ALT_FILL

    # Auto-fit column widths
    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


# ─── Word ─────────────────────────────────────────────────────────────────────


def _set_cell_text(cell, text: str) -> None:
    """
    Write text into a Word table cell by directly manipulating XML.

    Both `cell.text = value` (destroys style XML) and the paragraph/run API
    (unreliable on newly-cloned rows) fail on complex styled documents.
    Direct XML manipulation is the only approach that works reliably.
    """
    tc = cell._tc
    # Find the first <w:t> element; set its text directly
    t_elems = tc.findall(f".//{{{_W}}}t")
    if t_elems:
        t_elems[0].text = text
        # Preserve whitespace for values that start/end with spaces
        if text and (text[0] == " " or text[-1] == " "):
            t_elems[0].set(
                "{http://www.w3.org/XML/1998/namespace}space", "preserve"
            )
    else:
        # Cell has no run at all — build minimal paragraph/run/text XML
        p = tc.find(f"{{{_W}}}p")
        if p is None:
            from lxml import etree
            p = etree.SubElement(tc, f"{{{_W}}}p")
        r = p.find(f"{{{_W}}}r")
        if r is None:
            from lxml import etree
            r = etree.SubElement(p, f"{{{_W}}}r")
        t = r.find(f"{{{_W}}}t")
        if t is None:
            from lxml import etree
            t = etree.SubElement(r, f"{{{_W}}}t")
        t.text = text


def _clone_row(table, row_data: dict, col_map: dict[int, str]) -> None:
    """
    Add a new data row by deep-copying the last existing row's XML, then
    stamping in field values.  Cloning preserves every style, border, spacing
    and font attribute from the template row — table.add_row() does not.
    """
    # Use the last row as a style template (avoids cloning the header)
    template_tr = table.rows[-1]._tr
    new_tr = deepcopy(template_tr)

    # Blank out ALL text nodes in the clone first
    for t_el in new_tr.findall(f".//{{{_W}}}t"):
        t_el.text = ""

    # Attach to the table XML
    table._tbl.append(new_tr)
    new_row = table.rows[-1]

    # Write each field value into the matching cell
    for col_idx, field in col_map.items():
        if col_idx < len(new_row.cells):
            value = str(row_data.get(field, "") or "")
            _set_cell_text(new_row.cells[col_idx], value)


def _find_best_word_table(doc, columns: list[str]):
    """
    Return the table whose header row best matches the agent column names.
    Also returns the score so callers can decide whether to create a new table.
    """
    best_table = None
    best_score = -1

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if not any(headers):
            continue
        score = sum(
            max((_match_score(col, h) for h in headers), default=0)
            for col in columns
        )
        if score > best_score:
            best_score = score
            best_table = table

    return best_table, best_score


def _populate_word(file_bytes: bytes, agent_result: dict) -> bytes:
    data: list[dict] = agent_result.get("data", [])
    columns: list[str] = agent_result.get("column_order", []) or list(data[0].keys())
    summary: str = agent_result.get("summary", "")

    if not data:
        raise ValueError("No data was extracted from Gmail.")

    doc = Document(BytesIO(file_bytes))

    table, score = _find_best_word_table(doc, columns)

    if table and score > 0:
        # ── Append rows to the matched table ──────────────────────────────
        doc_headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Build col_index → agent field mapping using fuzzy matching
        col_map: dict[int, str] = {}
        for col_idx, header in enumerate(doc_headers):
            field = _best_col_for_header(header, columns)
            if field:
                col_map[col_idx] = field

        # Log the mapping so we can diagnose mismatches in the server log
        print(f"[doc_service] doc_headers={doc_headers}")
        print(f"[doc_service] columns={columns}")
        print(f"[doc_service] col_map={col_map}")
        print(f"[doc_service] data[0]={data[0] if data else 'empty'}")

        for row_data in data:
            # Clone last row instead of add_row() — preserves all style XML
            _clone_row(table, row_data, col_map)

    else:
        # ── No matching table found — create a new one at the end ─────────
        doc.add_heading("Gmail Data Export", level=1)
        doc.add_paragraph(summary).runs[0].italic = True
        doc.add_paragraph("")

        new_table = doc.add_table(rows=1, cols=len(columns))
        new_table.style = "Table Grid"

        # Header row
        for col_idx, col_name in enumerate(columns):
            cell = new_table.rows[0].cells[col_idx]
            _set_cell_text(cell, col_name, bold=True)

        # Data rows
        for row_data in data:
            new_row = new_table.add_row()
            for col_idx, field in enumerate(columns):
                if col_idx < len(new_row.cells):
                    value = str(row_data.get(field, "") or "")
                    _set_cell_text(new_row.cells[col_idx], value)

    # Footer note (always added)
    doc.add_paragraph("")
    footer_para = doc.add_paragraph(f"Updated by Gmail Helper — {summary}")
    if footer_para.runs:
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        footer_para.runs[0].italic = True

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
